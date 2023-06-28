import re
import os
import requests
import json
import openai
from bs4 import BeautifulSoup
import time
import random
from googleapiclient.discovery import build
import io
from PyPDF2 import PdfReader
import markdown
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from html2text import html2text
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

openai.api_key="openai-api-key"

def search_web(query, num_results=9):
    global skip
    search_results = []
    service = build("customsearch", "v1", developerKey="google-api-key")
    res = service.cse().list(q=query, cx='search-engine-key', num=num_results).execute()
    try:
        search_results = [item['link'] for item in res['items']]
        skip = False
    except:
        skip = True
        search_results = []
        while True:
            newlink = input("Search failed, enter links one by one manually and write DONE when done: ")
            if newlink == "DONE":
                break
            else:
                search_results.append(newlink)
    return search_results

def markdown_to_word(markdown_string, file_name):
    # Split the markdown text into paragraphs
    paragraphs = markdown_string.split('\n')
    # Create a new Word document
    doc = Document()
    # Iterate through each paragraph in the markdown text
    for para in paragraphs:
        if para.startswith("# "):
            # This is a heading 1
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para[2:])
            run.bold = True
            run.font.size = Pt(12)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif para.startswith("## "):
            # This is a heading 2
            paragraph = doc.add_paragraph()
            run = paragraph.add_run(para[3:])
            run.bold = True
            run.font.size = Pt(11)
        else:
            # This is a normal paragraph
            paragraph = doc.add_paragraph()
            # Use a regular expression to find and replace all instances of [^x^]
            superscript_pattern = re.compile(r'\[\^(\d+)\^\]')
            start = 0
            for match in superscript_pattern.finditer(para):
                # Add the text before the superscript as normal text
                run = paragraph.add_run(para[start:match.start()])
                # Add the superscript text
                run = paragraph.add_run(match.group(1))
                run.font.superscript = True
                start = match.end()
            # Add any remaining text after the last superscript
            run = paragraph.add_run(para[start:])
    # Save the document
    doc.save(file_name)
    # Open the document using the default application
    os.system(f'start {file_name}')

def extract_text_from_link(link: str):
    if link.endswith(('.htm')):
        page = requests.get(link)
        time.sleep(1)
        soup = BeautifulSoup(page.content, 'html.parser')
        all_b_tags = soup.find_all('b')
        return '\n'.join([b.get_text() for b in all_b_tags])
    elif link.endswith('.pdf'):
        response = requests.get(link)
        time.sleep(1)
        pdf_file = io.BytesIO(response.content)
        pdf_reader = PdfReader(pdf_file)
        text = ''
        for page in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page].extract_text()
        return text
    elif 'sec.gov' in link:
        page = requests.get(link)
        time.sleep(1)
        soup = BeautifulSoup(page.content, 'html.parser')
        all_font_tags = soup.find_all('font')
        return '\n'.join([font.get_text() for font in all_font_tags])
    else:
        page = requests.get(link)
        time.sleep(1)
        soup = BeautifulSoup(page.content, 'html.parser')
        all_p_tags = soup.find_all('p')
        return '\n'.join([p.get_text() for p in all_p_tags])


def passage_segmenter(passage):
    segment = []
    count = 0
    while count < len(passage):
        segment.append(passage[count:count + 12000])
        count += 12000
    return segment

def ask_question(messages):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=messages,
        stream=True
    )
    output = ""
    print("thinking...")
    for chunk in response:
        if "delta" in chunk["choices"][0]:
            delta = chunk["choices"][0]["delta"]
            if "content" in delta:
                content = delta["content"]
                output += content
                print(content, end="")

    return output

def order_links(query, links_str):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a helpful, pattern-following assistant."},
            {"role": "user", "content": f"Output in the format of a python list the order in which of following links are most likely to best answer the query, {query}\nExample output formatting (order is random. Use as a formatting example only. You are only allowed to re-oder the list, you are not allowed to remove links unless they are PDFs, which absolutely must be excluded from the list, or you have been told to leave them out): [4,7,2,1,5,6,8,9,3]\n (answer with list of ints only)? Links: " + links_str}
        ]
    )
    output_str = response["choices"][0]["message"]["content"].lower().replace(" ", "")
    # Process the output string and convert it into a list of integers
    output_list = [int(x) for x in output_str.strip('[]').split(',') if x.strip().isdigit()]
    return output_list

def name_file(output):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": """You are WordDocumentNameGeneratorGPT. You are given a research report and you come up with a name for it based on what it is about and based on the following word file naming conventions: Avoid backslashes in actual file or directory names as they serve as separators. Volume names may require backslashes, like in "C: or "server\share" in UNC names. Don't count on case sensitivity; treat OSCAR, Oscar, oscar as identical, although POSIX systems may differ. NTFS has POSIX case sensitivity support, but it's non-default. Unicode, and characters from the extent 128-255 are allowed barring reserved characters (like <, >, :, ", /, |, ?, *, and ASCII NUL) and integer values 1-31. Reserved file names include CON, PRN, AUX, NUL, COM(0-9), and LPT(0-9) including any extensions (e.g., NUL.txt). Also, avoid ending names with a space or a period. A period can represent the current directory (.\temp.txt) or the preceding directory (..\temp.txt). Note: Your entire output will be treated as the input for the file name so do not output anything you do not intend to be placed directly into the file name."""},
            {"role": "user", "content": f"Come up with a file name for the following research report (output file name only): {output}"}
        ]
    )
    output_str = response["choices"][0]["message"]["content"]
    return output_str

def summarize(query, res, link):
    response = openai.ChatCompletion.create(
        model="gpt-4-0613",
        messages=[
            {"role": "system", "content": """You are a helpful, pattern-following assistant. You are given some text retrieved from a website and a research query and you generate a very detailed and comprehensive summary of only the parts of the text relevant and useful to the answering the research query. Include as much detail as is physically possible. You only answer using the following JSON format and strictly follow JSON formatting conventions:
             {
                "is_relevant" : boolean, #true if the provided text provides information relevant to answering the users query, false if the text irrelevant to the query or just discusses access denial to a webpage
                "summary": "string" #very detailed (but not wordy) summary of the key information in the text if is_relevant is true (squeeze as much info into as few characters as you can. If you'd like, you can even use some shorthand). null if is_relevant is false
            }
            NOTE: It cannot be stressed enough how important it is that you do not break JSON formatting conventions. If you do, it will cause a JSONDecodeError"""},
            {"role": "user", "content": f"Summarize the key information in the following text in significant detail, which was scraped from the website {link}, that are relevant to answering the question, {query}. Text: " + res}
        ],
        stream = True
    )
    output = ""
    for chunk in response:
        if "delta" in chunk["choices"][0]:
            delta = chunk["choices"][0]["delta"]
            if "content" in delta:
                content = delta["content"]
                output += content
                print(content, end="")
    return output.strip()

def check_source(query_links, search_info, topic_summary):
    print("\n")
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": """You are a helpful, pattern-following assistant. You will be given a summary of the text from a website and a users question and you will decide whether the summary generated provides enough informration to answer the question or if the information gathered is inadequate and more is needed. You only respond in the following JSON format:
             {
                "continue" : boolean, #false if the provided summary provides ample information to answer the users question, true if more information is needed.
            }

            *Note: only set continue to true if key information is needed, otherwise set this value to false. If the question can be answered using the information gathered, then continue should be false. All thats needed is a few sentences that do answer the question.
            """},
            {"role": "user", "content": f"Does the following text provide ample information to answer the question, {search_info} Text: " + topic_summary +"\nLinks: " + str(query_links)}
        ],
        temperature=0,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
        stream = True
    )
    output = ""
    for chunk in response:
        if "delta" in chunk["choices"][0]:
            delta = chunk["choices"][0]["delta"]
            if "content" in delta:
                content = delta["content"]
                output += content
                print(content, end="")
    return output

def start_research(query):
    messages=[
        {"role": "system", "content": """
    You are an AI research assistant that only responds in JSON. You have been shown to be capable of completing
    complex research tasks that require retriving information on several different topics at a superhuman level.
    The following is the JSON format of all your outputs
    {
        "number_of_searches" : int #An integer value representing the number individual google searches needed,
        "search_queries" : {
            "1" : "string", #first search query,
            "2" : "string", #second search query
            ...
            "n" : "string" #nth query (Max queries: 8)
        }
        "search_query_goals" : {
            "1" : "string", #Describe what information that needs to be obtained with query one
            "2" : "string", #Describe what information that needs to be obtained with query two
            ...
            "n" : "string" #Describe what information that needs to be obtained with query n (Max queries: 5)
        }
    }

    search_query_goal should always be phrased as a question
    """}
    ]
    messages.append({"role": "user", "content": query})
    # Create a dictionary to hold search queries
    json_dict = {
        "search_queries": {}
    }
    response = ask_question(messages)
    json_data = json.loads(response)
    num_searches = int(json_data['number_of_searches'])
    return json_dict, json_data, num_searches

# Function to gather web information based on a search query
def get_web_info(json_dict, json_data, i, read_links):
    # Create a new key for the dictionary based on the index of the topic
    new_key = f"topic_{str(i)}"
    # Initialize an empty list to store search queries for the new topic
    json_dict["search_queries"][new_key] = []
    # Get the search query and goal from the input data
    search_query = json_data['search_queries'][str(i)]
    search_info = json_data['search_query_goals'][str(i)]
    # Print a status update
    print("\nRetrieving links for topic " + str(i) + "...\n")
    # Perform a web search based on the search query and store the result links
    query_links = search_web(search_query)
    # Remove any links that have already been read
    for link in read_links:
        if link in query_links:
            query_links.remove(link)
    # Print the retrieved links
    print("\nLinks retrieved:")
    for link_count, link in enumerate(query_links, 1):
        print(f"Link #{link_count}: {link}")
    # Order the links based on their relevancy to the search query
    print("\nOrdering links by relevancy...")
    ordered_links = order_links(search_info, str(query_links))
    print(f"Order chosen: {ordered_links}")
    # Return the search query, search goal, query links, ordered links, updated dictionary, and the new key
    return search_query, search_info, query_links, ordered_links, json_dict, new_key

# Function to create summaries from the information gathered from the web
def create_summaries(search_query, search_info, query_links, ordered_links, json_dict, json_data, new_key, i):
    # Initialize a dictionary to store whether to continue searching
    link_verdict_json = {
        "continue": True
    }
    # Initialize a link counter
    link_count = 1
    # Initialize flags for whether all links have been processed and whether a new search has been generated
    all_links_processed = False
    new_search_generated = False
    # Continue to process links as long as there are links left and the 'continue' flag is set
    while link_verdict_json['continue'] and not all_links_processed:
        # For each ordered link...
        for ordered_link in ordered_links:
            # Generate a new key for the link
            link_key = f"link_{str(link_count)}"
            # Get the link and its text content
            current_link = query_links[ordered_link - 1]
            current_link_text = extract_text_from_link(current_link).strip()
            print("\nSummarizing text...\n")
            print("Link: " + str(current_link) + "\n")
            # Segment the text content into manageable chunks
            segments = passage_segmenter(current_link_text)
            # For each segment, generate a summary and check if the source is relevant
            for segment in segments:
                link_summary = summarize(search_query, segment, current_link)
                summary_json = json.loads(link_summary)
                # If the source is relevant, add it to the dictionary and check if more information is needed
                if summary_json['is_relevant']:
                    read_links.append(current_link)
                    new_value_dict = {
                        f"{link_key}": current_link,
                        f"{link_key}_summary": summary_json['summary']
                    }
                    json_dict["search_queries"][new_key].append(new_value_dict)
                    link_verdict_json = json.loads(check_source(query_links, search_info, summary_json['summary']))
                    # Break the loop if no more information is needed or if the length limit has been reached
                    if not skip:
                        if not link_verdict_json['continue'] or len(link_summary) >= 12000:
                            break
                else:
                    break
                print("\n")
            # Increment the link counter
            link_count += 1
            # Set the 'all_links_processed' flag if all links have been processed
            if link_count > len(ordered_links):
                all_links_processed = True
            # Break the loop if no more information is needed, if the link limit has been reached, or if all links have been processed
            if not link_verdict_json['continue'] or link_count > 9 or all_links_processed:
                break
        if not link_verdict_json['continue'] or link_count > 9 or all_links_processed:
            break
    # Print the updated dictionary
    storage = json.dumps(json_dict, indent=2)
    print("\n")
    print(storage)
    # Return the updated dictionary and the 'new_search_generated' flag
    return json_dict, new_search_generated

# Function to generate an answer from the gathered and summarized information
def generate_answer(query, json_dict):
    # Initialize the list of messages for the answer
    answer_messages=[
        {"role": "system", "content": """
        You are a research chatbot. You will be provided with a research task from the user as well as a bunch of information that was just scraped from the web and your job is to use that information to generate a very detailed and comprehensive research report with evidence-based explanations for every argument. Your reports should be comparable in length to professional industry research reports like ones published by Nielsens or think tanks like Brookings Institute. Your outputs should never be less than 2000 words in length but you should always aim for 3200 words. You will always cite your work by using footnotes and your output will always be in markdown syntax.
        Follow this formatting:
            ## Main title
            # Sub Titles
            [^n^] (where n is a number) for footnotes
        """}
    ]
    # Generate the prompt for the answer
    answer_prompt = "User generated research question: " + query + "\n\nInformation:\n"
    sorted_keys = sorted(json_dict["search_queries"].keys())
    # For each topic, add its information to the prompt
    for topic in sorted_keys:
        answer_prompt += f"\n- Topic: {topic}\n"
        for topic_info_dict in json_dict["search_queries"][topic]:
            for link_key, link in topic_info_dict.items():
                # Only add the link to the prompt if it does not contain a summary
                if "_summary" not in link_key:
                    answer_prompt += f"  - {link_key}: {link}\n"
                    summary_key = f"{link_key}_summary"
                    # If a summary exists for the link, add it to the prompt
                    if summary_key in topic_info_dict:
                        summary = topic_info_dict[summary_key]
                        answer_prompt += f"  - {summary_key}: {summary}\n"
    # Print the prompt and add it to the messages
    print(answer_prompt)
    answer_messages.append({"role": "user", "content": answer_prompt})
    # Generate the answer based on the messages
    answer = ask_question(answer_messages)
    # Return the answer
    return answer


def main():
    query = input("Query: ")
    global read_links
    read_links = []
    json_dict, json_data, num_searches = start_research(query)
    for i in range(1, num_searches + 1):
        search_query, search_info, query_links, ordered_links, json_dict, new_key = get_web_info(json_dict, json_data, i, read_links)
        json_dict, new_search_generated = create_summaries(search_query, search_info, query_links, ordered_links, json_dict, json_data, new_key, i)
    output = generate_answer(query, json_dict)
    file_name = name_file(output).replace(" ", "_")
    markdown_to_word(output, file_name)

main()
